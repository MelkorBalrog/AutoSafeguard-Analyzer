"""Simple WYSIWYG requirements document editor."""

from __future__ import annotations

import tkinter as tk
from tkinter import ttk, filedialog, simpledialog, font
import webbrowser
from typing import Dict, Iterable

try:
    from PIL import Image, ImageTk
except Exception:  # pragma: no cover - optional dependency
    Image = ImageTk = None

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches
except Exception:  # pragma: no cover - optional dependency
    Document = None
    OxmlElement = qn = RELATIONSHIP_TYPE = Inches = None


class RequirementsDocumentEditor(tk.Toplevel):  # pragma: no cover - GUI
    """A minimal WYSIWYG editor for requirement specifications.

    The editor provides heading styles, font controls and allows inserting
    images and hyperlinks. Content can be persisted into a ``.docx`` file
    using :mod:`python-docx` for later export. When *requirements* are
    provided, they are inserted into the document grouped under a heading for
    ``req_type``.
    """

    def __init__(
        self,
        master: tk.Misc | None = None,
        *,
        req_type: str | None = None,
        requirements: Iterable[dict] | None = None,
    ) -> None:
        super().__init__(master)
        title = f"{req_type.title()} Specification" if req_type else "Requirements Specification"
        self.title(title)
        self.geometry("900x700")
        self.images: Dict[str, tuple] = {}
        self.hyperlinks: Dict[str, str] = {}
        self._link_counter = 0

        self._create_toolbar()
        self._create_text_widget()

        if req_type and requirements:
            self._populate_requirements(req_type, requirements)

    # ------------------------------------------------------------------
    def _populate_requirements(
        self, req_type: str, requirements: Iterable[dict]
    ) -> None:
        """Insert *requirements* grouped under *req_type* heading."""
        self.text.insert("end", f"{req_type.title()} Requirements\n", ("heading1",))
        for req in requirements:
            rid = req.get("id", "")
            txt = req.get("text", "")
            self.text.insert("end", f"{rid}: {txt}\n")

    # ------------------------------------------------------------------
    def _create_toolbar(self) -> None:
        bar = ttk.Frame(self)
        bar.pack(fill=tk.X)

        ttk.Button(bar, text="H1", command=lambda: self._apply_tag("heading1")).pack(side=tk.LEFT)
        ttk.Button(bar, text="H2", command=lambda: self._apply_tag("heading2")).pack(side=tk.LEFT)
        ttk.Button(bar, text="Bold", command=lambda: self._apply_tag("bold")).pack(side=tk.LEFT)
        ttk.Button(bar, text="Italic", command=lambda: self._apply_tag("italic")).pack(side=tk.LEFT)
        ttk.Button(bar, text="Underline", command=lambda: self._apply_tag("underline")).pack(side=tk.LEFT)

        fonts = sorted(set(font.families()))
        self.font_var = tk.StringVar(value="TkDefaultFont")
        self.size_var = tk.StringVar(value="12")
        ttk.Combobox(bar, values=fonts, textvariable=self.font_var, width=15).pack(side=tk.LEFT, padx=5)
        ttk.Combobox(bar, values=["8", "9", "10", "11", "12", "14", "16", "18", "20", "24"], textvariable=self.size_var, width=3).pack(side=tk.LEFT)
        ttk.Button(bar, text="Apply Font", command=self._apply_font).pack(side=tk.LEFT, padx=5)

        ttk.Button(bar, text="Insert Image", command=self._insert_image).pack(side=tk.LEFT, padx=5)
        ttk.Button(bar, text="Insert Link", command=self._insert_link).pack(side=tk.LEFT)
        ttk.Button(bar, text="Save", command=self._save_docx).pack(side=tk.RIGHT)

    # ------------------------------------------------------------------
    def _create_text_widget(self) -> None:
        self.text = tk.Text(self, wrap="word")
        self.text.pack(fill=tk.BOTH, expand=True)

        base = font.Font(font=self.text.cget("font"))
        h1 = font.Font(font=base)
        h1.configure(size=18, weight="bold")
        h2 = font.Font(font=base)
        h2.configure(size=16, weight="bold")
        bold = font.Font(font=base)
        bold.configure(weight="bold")
        italic = font.Font(font=base)
        italic.configure(slant="italic")
        underline = font.Font(font=base)
        underline.configure(underline=True)

        self.text.tag_configure("heading1", font=h1)
        self.text.tag_configure("heading2", font=h2)
        self.text.tag_configure("bold", font=bold)
        self.text.tag_configure("italic", font=italic)
        self.text.tag_configure("underline", font=underline)

    # ------------------------------------------------------------------
    def _apply_tag(self, tag: str) -> None:
        try:
            self.text.tag_add(tag, "sel.first", "sel.last")
        except tk.TclError:
            pass

    # ------------------------------------------------------------------
    def _apply_font(self) -> None:
        try:
            font_name = self.font_var.get()
            size = int(self.size_var.get())
            f = font.Font(family=font_name, size=size)
            tag = f"font_{font_name}_{size}"
            if not self.text.tag_cget(tag, "font"):
                self.text.tag_configure(tag, font=f)
            self.text.tag_add(tag, "sel.first", "sel.last")
        except Exception:
            pass

    # ------------------------------------------------------------------
    def _insert_image(self) -> None:
        if Image is None:
            return
        path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png *.jpg *.jpeg *.gif")])
        if not path:
            return
        img = Image.open(path)
        tk_img = ImageTk.PhotoImage(img)
        img_id = self.text.image_create("insert", image=tk_img)
        self.images[img_id] = (tk_img, path)
        self.text.insert("insert", "\n")

    # ------------------------------------------------------------------
    def _insert_link(self) -> None:
        try:
            sel = self.text.get("sel.first", "sel.last")
        except tk.TclError:
            return
        url = simpledialog.askstring("Hyperlink", "Enter URL:")
        if not url:
            return
        tag = f"link_{self._link_counter}"
        self._link_counter += 1
        self.hyperlinks[tag] = url
        self.text.tag_add(tag, "sel.first", "sel.last")
        self.text.tag_configure(tag, foreground="blue", underline=True)
        self.text.tag_bind(tag, "<Button-1>", lambda e, u=url: webbrowser.open_new(u))

    # ------------------------------------------------------------------
    def _save_docx(self) -> None:
        if Document is None:
            return
        fname = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word Document", "*.docx")]
        )
        if not fname:
            return
        doc = Document()
        dump = self.text.dump("1.0", "end-1c", text=True, tag=True, image=True)
        p = doc.add_paragraph()
        active_tags: set[str] = set()
        current_link: str | None = None
        for kind, value, index in dump:
            if kind == "text":
                run = None
                if current_link:
                    run = self._add_hyperlink(p, current_link, value)
                else:
                    run = p.add_run(value)
                run.bold = "bold" in active_tags
                run.italic = "italic" in active_tags
                run.underline = "underline" in active_tags
                if "heading1" in active_tags or "heading2" in active_tags:
                    pass
            elif kind == "image":
                img = self.images.get(value)
                if img:
                    doc.add_picture(img[1], width=Inches(4))
                    p = doc.add_paragraph()
            elif kind == "tagon":
                active_tags.add(value)
                if value in ("heading1", "heading2"):
                    p = doc.add_paragraph()
                    p.style = "Heading 1" if value == "heading1" else "Heading 2"
                if value.startswith("link_"):
                    current_link = self.hyperlinks.get(value)
            elif kind == "tagoff":
                if value in active_tags:
                    active_tags.remove(value)
                if value.startswith("link_"):
                    current_link = None
        doc.save(fname)

    # ------------------------------------------------------------------
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Create a hyperlink within *paragraph* pointing to *url*."""
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        rPr.append(color)
        new_run.append(rPr)
        text_elem = OxmlElement("w:t")
        text_elem.text = text
        new_run.append(text_elem)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return paragraph.add_run()
